import docx

doc = docx.Document()
doc.add_paragraph('Hello World!')
paraObj1 = doc.add_paragraph('This is a second Paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another Paragraph.')
paraObj1.add_run('This text is being added to the second Paragraph.')
doc.add_paragraph('Tunde Muhamed', 'Title')
doc.save('multipleparagraphs.docx')