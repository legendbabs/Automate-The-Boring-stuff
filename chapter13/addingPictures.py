import docx

doc = docx.Document()
doc.add_heading('Heading 1', 1)
doc.add_heading('Heading 2', 2)
doc.add_heading('Heading 3', 3)
doc.add_heading('Heading 4', 4)
doc.add_picture('doc.jpg')
doc.save('myPicture.docx')