import docx

doc = docx.Document()
doc.add_heading('Salami Tunde Muhamed', 1)
doc.add_heading('Olajire Toyosi Kizran', 2)
doc.add_heading('Ajayi Elijah Taiwo', 3)
doc.add_heading('Oloruntele Nurudeen Atanda', 4)
doc.save('myNameHeadings.docx')