import docx

doc = docx.Document()
doc.add_heading('Heading 1', 1)
doc.add_heading('Heading 2', 2)
doc.add_heading('Heading 3', 3)
doc.add_heading('Heading 4', 4)
doc.add_picture('doc.jpg', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save('myPicture2.docx')